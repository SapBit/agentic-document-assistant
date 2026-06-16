"""
File handling utilities for uploading and extracting text from different file formats.
Supports PDF, DOCX, and TXT files.
"""

import PyPDF2
from docx import Document
import os
from datetime import datetime


def extract_text_from_pdf(file_path):
    """
    Extract text from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
    """
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text()
        return text
    except Exception as e:
        return f"Error extracting PDF: {str(e)}"


def extract_text_from_docx(file_path):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        return f"Error extracting DOCX: {str(e)}"


def extract_text_from_txt(file_path):
    """
    Extract text from a TXT file.
    
    Args:
        file_path (str): Path to the TXT file
        
    Returns:
        str: Text content from the file
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        return f"Error extracting TXT: {str(e)}"


def extract_text_from_file(file_path):
    """
    Extract text from any supported file format.
    Automatically detects file type based on extension.
    
    Args:
        file_path (str): Path to the file
        
    Returns:
        str: Extracted text content
    """
    file_ext = os.path.splitext(file_path)[1].lower()
    
    if file_ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        return extract_text_from_docx(file_path)
    elif file_ext == '.txt':
        return extract_text_from_txt(file_path)
    else:
        return f"Unsupported file format: {file_ext}"


def save_output(content, filename_prefix, file_format='txt'):
    """
    Save processed content to a file in the outputs folder.
    
    Args:
        content (str): Content to save
        filename_prefix (str): Prefix for the filename
        file_format (str): File format (txt, docx, or pdf)
        
    Returns:
        str: Path to the saved file
    """
    os.makedirs('outputs', exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"outputs/{filename_prefix}_{timestamp}.{file_format}"
    
    try:
        if file_format == 'docx':
            doc = Document()
            doc.add_paragraph(content)
            doc.save(filename)
        else:  # txt format
            with open(filename, 'w', encoding='utf-8') as f:
                f.write(content)
        return filename
    except Exception as e:
        return f"Error saving file: {str(e)}"
